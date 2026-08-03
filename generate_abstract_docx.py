import docx
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

doc = docx.Document()

# Set standard margins (1 inch)
for section in doc.sections:
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)

# Default style: Times New Roman, 12pt
style = doc.styles['Normal']
style.font.name = 'Times New Roman'
style.font.size = Pt(12)

def add_p(text, align=WD_ALIGN_PARAGRAPH.LEFT, bold=False, italic=False, space_after=6, line_spacing=1.5):
    p = doc.add_paragraph()
    p.alignment = align
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.line_spacing = line_spacing
    run = p.add_run(text)
    run.font.name = 'Times New Roman'
    run.font.size = Pt(12)
    run.bold = bold
    run.italic = italic
    return p

# Header block matching HITAM template
add_p('HYDERABAD INSTITUTE OF TECHNOLOGY AND MANAGEMENT', WD_ALIGN_PARAGRAPH.CENTER, bold=True, space_after=2)
add_p('DEPARTMENT OF COMPUTER SCIENCE AND ENGINEERING', WD_ALIGN_PARAGRAPH.CENTER, bold=True, space_after=2)
add_p('IV B. Tech-I Semester', WD_ALIGN_PARAGRAPH.CENTER, bold=True, space_after=2)
add_p('Mini Project A.Y 2026-27', WD_ALIGN_PARAGRAPH.CENTER, bold=True, space_after=12)

add_p('ABSTRACT', WD_ALIGN_PARAGRAPH.CENTER, bold=True, space_after=12)

# Metadata Section
p_meta = doc.add_paragraph()
p_meta.paragraph_format.line_spacing = 1.5
p_meta.paragraph_format.space_after = Pt(12)

r1 = p_meta.add_run('Team No: ')
r1.bold = True
r1.font.name = 'Times New Roman'
r1.font.size = Pt(12)

r2 = p_meta.add_run('[Fill Team No.]\n')
r2.font.name = 'Times New Roman'
r2.font.size = Pt(12)

r3 = p_meta.add_run('Project Title: ')
r3.bold = True
r3.font.name = 'Times New Roman'
r3.font.size = Pt(12)

r4 = p_meta.add_run('Learniverse AI: Dual-Persona RAG Educational Platform & Automated Placement Assessment System\n')
r4.font.name = 'Times New Roman'
r4.font.size = Pt(12)

r5 = p_meta.add_run('Abstract:')
r5.bold = True
r5.font.name = 'Times New Roman'
r5.font.size = Pt(12)

# Condensed Abstract Paragraphs (~180 words, Times New Roman, 12pt, 1.5 line spacing, Justified)
p1 = add_p(
    'Traditional online learning platforms and generic AI chatbots often fail to deliver reliable computer science education without generating inaccurate answers ("hallucinations") or dense jargon. To address this, Learniverse AI is an interactive learning and placement assessment platform designed for Computer Science students mastering Data Structures & Algorithms (DSA) and technical placement preparation.',
    align=WD_ALIGN_PARAGRAPH.JUSTIFY,
    space_after=6,
    line_spacing=1.5
)

p2 = add_p(
    'The core architecture features a Dual-Persona Retrieval-Augmented Generation (RAG) pipeline. By indexing academic textbooks into a persistent vector database (ChromaDB) using Google Gemini Embeddings, Learniverse AI grounds all answers in verified textbook content with direct chapter citations, eliminating AI hallucinations. It provides dual AI mentors: a Teacher AI for formal academic explanations and a Peer AI for intuitive, analogy-based learning.',
    align=WD_ALIGN_PARAGRAPH.JUSTIFY,
    space_after=6,
    line_spacing=1.5
)

p3 = add_p(
    'Additionally, Learniverse AI incorporates an automated Placement Assessment System evaluating students across Quantitative Aptitude, Logical Reasoning, Verbal Ability, and Coding. Built with React, Vite, and TailwindCSS on the frontend and FastAPI, Python, ChromaDB, and Google Gemini API on the backend, Learniverse AI effectively bridges conceptual learning with placement readiness.',
    align=WD_ALIGN_PARAGRAPH.JUSTIFY,
    space_after=14,
    line_spacing=1.5
)

# Keywords section
p_kw = doc.add_paragraph()
p_kw.paragraph_format.line_spacing = 1.5
p_kw.paragraph_format.space_after = Pt(28)

r_kw_title = p_kw.add_run('Keywords: ')
r_kw_title.bold = True
r_kw_title.font.name = 'Times New Roman'
r_kw_title.font.size = Pt(12)

r_kw_body = p_kw.add_run('Retrieval-Augmented Generation (RAG), AI in Education, Dual-Persona LLM, Data Structures & Algorithms (DSA), ChromaDB, Gemini API, Placement Assessment System.')
r_kw_body.italic = True
r_kw_body.font.name = 'Times New Roman'
r_kw_body.font.size = Pt(12)

# Signatures Block
p_sig = doc.add_paragraph()
p_sig.paragraph_format.line_spacing = 1.5
p_sig.paragraph_format.space_after = Pt(28)
r_sig = p_sig.add_run('Project Guide                 Project Coordinator                 HOD, CSE')
r_sig.bold = True
r_sig.font.name = 'Times New Roman'
r_sig.font.size = Pt(12)

# Team Details Block
p_team = doc.add_paragraph()
p_team.paragraph_format.line_spacing = 1.5
r_t_title = p_team.add_run('Team Details:\n')
r_t_title.bold = True
r_t_title.font.name = 'Times New Roman'
r_t_title.font.size = Pt(12)

members = [
    'Member 1: Name (Roll No.)',
    'Member 2: Name (Roll No.)',
    'Member 3: Name (Roll No.)',
    'Member 4: Name (Roll No.)'
]

for idx, m in enumerate(members):
    r_m = p_team.add_run(m + ('\n' if idx < len(members)-1 else ''))
    r_m.font.name = 'Times New Roman'
    r_m.font.size = Pt(12)

out_path = r'c:\Users\dhanu\Downloads\learniverse-ai-main (1)\learniverse-ai-main\backend\placement_assessment_system\Abstract_Learniverse_AI_HITAM.docx'
doc.save(out_path)
print(f'Saved updated docx successfully to {out_path}')
